#!/usr/bin/env python3
"""
Enhanced Flask Backend for Threat Modeling Pipeline with Review System
Includes quality checkpoints, confidence scoring, and collaborative review
"""

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from flask_socketio import SocketIO, emit
import subprocess
import json
import os
import sys
import tempfile
import shutil
from pathlib import Path
from datetime import datetime
import logging
import traceback
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
import threading
import time
import uuid
import re

# IMPORTANT: Load .env BEFORE anything else
load_dotenv()

# Import config function
def get_config():
    """Get configuration from environment with defaults."""
    return {
        'llm_provider': os.getenv('LLM_PROVIDER', 'scaleway'),
        'llm_model': os.getenv('LLM_MODEL', 'llama-3.3-70b-instruct'),
        'local_llm_endpoint': os.getenv('LOCAL_LLM_ENDPOINT', 'http://localhost:11434/api/generate'),
        'custom_system_prompt': os.getenv('CUSTOM_SYSTEM_PROMPT', ''),
        'timeout': int(os.getenv('PIPELINE_TIMEOUT', '5000')),
        'input_dir': os.getenv('INPUT_DIR', './input_documents'),
        'output_dir': os.getenv('OUTPUT_DIR', './output'),
        'dfd_output_path': os.getenv('DFD_OUTPUT_PATH', './output/dfd_components.json'),
        'mitre_enabled': os.getenv('MITRE_ENABLED', 'true').lower() == 'true',
        'mitre_version': os.getenv('MITRE_VERSION', 'v13.1'),
        'scw_api_url': os.getenv('SCW_API_URL', 'https://api.scaleway.ai/v1'),
        'scw_secret_key': os.getenv('SCW_SECRET_KEY') or os.getenv('SCW_API_KEY')
    }

# Document processing imports with error handling
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("Warning: PyPDF2 not available. PDF processing disabled.")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. DOCX processing disabled.")

# Global configuration (runtime overrides)
runtime_config = get_config()

app = Flask(__name__)
CORS(app)

# Initialize SocketIO
socketio = SocketIO(app, cors_allowed_origins="*")

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Configuration
UPLOAD_FOLDER = './uploads'
OUTPUT_FOLDER = runtime_config['output_dir']
INPUT_FOLDER = runtime_config['input_dir']
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'doc', 'docx'}

# Ensure directories exist
for folder in [UPLOAD_FOLDER, OUTPUT_FOLDER, INPUT_FOLDER]:
    os.makedirs(folder, exist_ok=True)

# Verify environment on startup
print("\n" + "="*60)
print("THREAT MODELING PIPELINE BACKEND WITH REVIEW SYSTEM")
print("="*60)
api_key = runtime_config['scw_secret_key']
if api_key:
    print(f"✓ API Key loaded: ***{api_key[-4:]}")
else:
    print("⚠️  WARNING: No API key found in environment!")
    print("  Please ensure your .env file contains SCW_SECRET_KEY=your_key_here")
print(f"Working directory: {os.getcwd()}")
print(f"Python: {sys.executable}")

# Check for required scripts
scripts = ['info_to_dfds.py', 'dfd_to_threats.py', 'improve_threat_quality.py', 'attack_path_analyzer.py']
for script in scripts:
    if os.path.exists(script):
        print(f"✓ Found: {script}")
    else:
        print(f"✗ Missing: {script}")

print("="*60)

# Enhanced pipeline state with review system
pipeline_state = {
    'current_session': None,
    'logs': [],
    'step_outputs': {},
    'validations': {},
    'review_queue': {},      # Items needing review per step
    'review_history': [],    # Audit trail of all reviews
    'quality_metrics': {}    # Track quality improvements
}

# Lock for thread-safe operations
state_lock = threading.Lock()

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_file(file_path):
    """Extract text content from various file formats."""
    file_ext = file_path.lower().split('.')[-1]
    text_content = ""
    
    try:
        if file_ext == 'txt':
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                        text_content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
                    
        elif file_ext == 'pdf' and PDF_AVAILABLE:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
                        
        elif file_ext == 'docx' and DOCX_AVAILABLE:
            doc = DocxDocument(file_path)
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + "\t"
                    text_content += "\n"
                    
        elif file_ext in ['doc', 'docx'] and not DOCX_AVAILABLE:
            return None, "DOCX format not supported. Please install python-docx or convert to TXT."
            
        elif file_ext == 'pdf' and not PDF_AVAILABLE:
            return None, "PDF format not supported. Please install PyPDF2 or convert to TXT."
            
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {str(e)}")
        return None, str(e)
    
    return text_content, None

def add_log(message, log_type='info'):
    """Add a log entry to the pipeline state."""
    with state_lock:
        log_entry = {
            'timestamp': datetime.now().isoformat(),
            'type': log_type,
            'message': message
        }
        pipeline_state['logs'].append(log_entry)
        
        # Keep only last 1000 logs
        if len(pipeline_state['logs']) > 1000:
            pipeline_state['logs'] = pipeline_state['logs'][-1000:]
    
    logger.info(f"[{log_type}] {message}")

# ==================== REVIEW SYSTEM FUNCTIONS ====================

def calculate_confidence(value, value_type):
    """Calculate confidence score for extracted values."""
    confidence = 0.5  # Base confidence
    
    if value_type == 'entity':
        # Higher confidence for common patterns
        common_entities = ['user', 'admin', 'customer', 'system', 'api', 'external']
        if any(entity in value.lower() for entity in common_entities):
            confidence += 0.3
        
        # Check for proper naming convention
        if re.match(r'^[A-Z][a-zA-Z0-9_]*$', value):
            confidence += 0.1
            
    elif value_type == 'asset':
        # Database patterns
        if any(db in value.lower() for db in ['db', 'database', 'store', 'cache']):
            confidence += 0.2
        
        # File storage patterns
        if any(fs in value.lower() for fs in ['file', 'storage', 'blob', 's3']):
            confidence += 0.2
            
    elif value_type == 'process':
        # Service patterns
        if any(svc in value.lower() for svc in ['service', 'server', 'api', 'gateway']):
            confidence += 0.3
            
    elif value_type == 'data_flow':
        # Complete data flows have higher confidence
        if all(key in value for key in ['source', 'destination', 'protocol']):
            confidence += 0.2
        if 'authentication_mechanism' in value and value['authentication_mechanism'] != 'Unknown':
            confidence += 0.2
            
    return min(confidence, 0.95)  # Cap at 95%

def infer_criticality_hint(asset_name):
    """Infer criticality based on asset name and type."""
    asset_lower = asset_name.lower()
    
    if any(critical in asset_lower for critical in ['payment', 'billing', 'credential', 'secret', 'key']):
        return "Likely Critical - handles sensitive financial or authentication data"
    elif any(high in asset_lower for high in ['user', 'customer', 'profile', 'personal']):
        return "Likely High - contains user PII data"
    elif any(medium in asset_lower for medium in ['log', 'cache', 'session', 'temp']):
        return "Likely Medium - temporary or derived data"
    else:
        return "Consider data sensitivity and business impact"

def infer_exposure_hint(asset_name):
    """Infer exposure level based on asset name."""
    asset_lower = asset_name.lower()
    
    if any(public in asset_lower for public in ['public', 'cdn', 'static', 'frontend']):
        return "Likely Internet-facing - public resources"
    elif any(dmz in asset_lower for dmz in ['api', 'gateway', 'proxy', 'load']):
        return "Likely DMZ - exposed but protected services"
    elif any(internal in asset_lower for internal in ['db', 'database', 'internal', 'private']):
        return "Likely Internal - should not be directly exposed"
    else:
        return "Consider network architecture and access patterns"

def group_similar_threats(threats):
    """Group threats that might be duplicates."""
    groups = []
    processed = set()
    
    for i, threat1 in enumerate(threats):
        if i in processed:
            continue
            
        group = [threat1]
        processed.add(i)
        
        for j, threat2 in enumerate(threats[i+1:], i+1):
            if j in processed:
                continue
                
            # Simple similarity check based on component and category
            if (threat1.get('component_name') == threat2.get('component_name') and
                threat1.get('stride_category') == threat2.get('stride_category')):
                
                # Check description similarity (simple word overlap)
                words1 = set(threat1.get('threat_description', '').lower().split())
                words2 = set(threat2.get('threat_description', '').lower().split())
                
                if len(words1 & words2) > len(words1) * 0.5:
                    group.append(threat2)
                    processed.add(j)
        
        if len(group) > 1:
            groups.append(group)
    
    return groups

def generate_review_items(step, step_data):
    """Generate review items with confidence scores."""
    items = []
    
    if step == 2:  # DFD Extraction
        dfd = step_data.get('dfd', {})
        
        # Review external entities
        for entity in dfd.get('external_entities', []):
            confidence = calculate_confidence(entity, 'entity')
            if confidence < 0.8:
                items.append({
                    'id': str(uuid.uuid4()),
                    'type': 'external_entity',
                    'value': entity,
                    'confidence': confidence,
                    'status': 'pending',
                    'questions': [
                        'Is this correctly identified as an external entity?',
                        'Should this be classified as a process instead?'
                    ],
                    'suggestions': []
                })
        
        # Review assets with criticality
        for asset in dfd.get('assets', []):
            items.append({
                'id': str(uuid.uuid4()),
                'type': 'asset',
                'value': asset,
                'confidence': 0.6,  # Always review assets for criticality
                'status': 'pending',
                'attributes_needed': {
                    'criticality': {
                        'question': 'What is the criticality level?',
                        'options': ['Critical', 'High', 'Medium', 'Low'],
                        'hint': infer_criticality_hint(asset)
                    },
                    'exposure': {
                        'question': 'What is the exposure level?',
                        'options': ['Internet-facing', 'DMZ', 'Internal', 'Isolated'],
                        'hint': infer_exposure_hint(asset)
                    },
                    'data_classification': {
                        'question': 'What type of data does this store?',
                        'options': ['PII', 'PHI', 'PCI', 'Confidential', 'Public'],
                        'multiple': True
                    }
                }
            })
        
        # Review data flows
        for flow in dfd.get('data_flows', []):
            confidence = calculate_confidence(flow, 'data_flow')
            if not flow.get('authentication_mechanism') or flow['authentication_mechanism'] == 'Unknown':
                items.append({
                    'id': str(uuid.uuid4()),
                    'type': 'data_flow',
                    'value': flow,
                    'confidence': confidence,
                    'status': 'pending',
                    'missing_fields': ['authentication_mechanism'],
                    'questions': [
                        'What authentication method is used for this data flow?',
                        'Is encryption in transit implemented?'
                    ]
                })
    
    elif step == 3:  # Threat Generation
        threats = step_data.get('threats', [])
        
        # Group similar threats for review
        threat_groups = group_similar_threats(threats)
        for group in threat_groups:
            if len(group) > 1:
                items.append({
                    'id': str(uuid.uuid4()),
                    'type': 'duplicate_threats',
                    'value': group,
                    'confidence': 0.4,
                    'status': 'pending',
                    'action_needed': 'merge_or_differentiate',
                    'questions': [
                        'Are these threats describing the same vulnerability?',
                        'Should they be merged or kept separate?'
                    ]
                })
        
        # Review high-impact threats
        for threat in threats:
            if threat.get('impact') == 'Critical' and threat.get('likelihood') == 'High':
                items.append({
                    'id': str(uuid.uuid4()),
                    'type': 'high_risk_threat',
                    'value': threat,
                    'confidence': 0.9,
                    'status': 'pending',
                    'validation_needed': True,
                    'questions': [
                        'Is this risk assessment accurate?',
                        'Are there compensating controls in place?'
                    ]
                })
    
    return items

def apply_review_corrections(step, item, corrections):
    """Apply expert corrections to the pipeline data."""
    step_data = pipeline_state['step_outputs'].get(step, {})
    
    if step == 2:  # DFD corrections
        dfd = step_data.get('dfd', {})
        
        if item['type'] == 'asset' and corrections:
            # Find and update the asset
            for i, asset in enumerate(dfd.get('assets', [])):
                if asset == item['value']:
                    # Enrich with metadata
                    if 'assets_metadata' not in dfd:
                        dfd['assets_metadata'] = {}
                    
                    dfd['assets_metadata'][asset] = {
                        'criticality': corrections.get('criticality'),
                        'exposure': corrections.get('exposure'),
                        'data_classification': corrections.get('data_classification'),
                        'reviewed': True,
                        'reviewer': item['review']['reviewer'],
                        'review_date': item['review']['timestamp']
                    }
                    break
        
        elif item['type'] == 'data_flow' and corrections:
            # Update the data flow
            for flow in dfd.get('data_flows', []):
                if flow == item['value']:
                    flow.update(corrections)
                    flow['reviewed'] = True
                    break
    
    # Save updated data
    save_step_data(step, step_data)

def save_step_data(step, data):
    """Save step data to file."""
    files = {
        2: 'dfd_components.json',
        3: 'identified_threats.json',
        4: 'refined_threats.json',
        5: 'attack_paths.json'
    }
    
    if step in files:
        file_path = os.path.join(OUTPUT_FOLDER, files[step])
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)

def count_pending_reviews():
    """Count total pending reviews across all steps."""
    count = 0
    for step_items in pipeline_state.get('review_queue', {}).values():
        count += len([item for item in step_items if item['status'] == 'pending'])
    return count

def calculate_quality_metrics():
    """Calculate quality metrics from review history."""
    metrics = {
        'total_reviews': len(pipeline_state['review_history']),
        'approval_rate': 0,
        'average_confidence_improvement': 0,
        'most_common_issues': {},
        'reviewer_stats': {}
    }
    
    if pipeline_state['review_history']:
        approvals = len([r for r in pipeline_state['review_history'] if r['decision'] == 'approve'])
        metrics['approval_rate'] = approvals / len(pipeline_state['review_history'])
        
        # Calculate reviewer statistics
        for review in pipeline_state['review_history']:
            reviewer = review['reviewer']
            if reviewer not in metrics['reviewer_stats']:
                metrics['reviewer_stats'][reviewer] = {'count': 0, 'approvals': 0}
            metrics['reviewer_stats'][reviewer]['count'] += 1
            if review['decision'] == 'approve':
                metrics['reviewer_stats'][reviewer]['approvals'] += 1
    
    return metrics

# ==================== REVIEW API ENDPOINTS ====================

@app.route('/api/review-queue/<int:step>', methods=['GET'])
def get_review_queue(step):
    """Get items needing review for a specific step."""
    try:
        with state_lock:
            step_data = pipeline_state.get('step_outputs', {}).get(step)
            if not step_data:
                return jsonify({'error': 'Step not completed yet'}), 404
            
            # Generate review items if not already in queue
            if step not in pipeline_state['review_queue']:
                review_items = generate_review_items(step, step_data)
                pipeline_state['review_queue'][step] = review_items
            else:
                review_items = pipeline_state['review_queue'][step]
            
            pending_items = [item for item in review_items if item['status'] == 'pending']
            
            return jsonify({
                'step': step,
                'items': pending_items,
                'total': len(review_items),
                'pending': len(pending_items)
            })
    except Exception as e:
        logger.error(f"Review queue error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/review-item/<item_id>', methods=['POST'])
def review_item(item_id):
    """Submit review for a specific item."""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        with state_lock:
            # Find the item in review queue
            for step, items in pipeline_state.get('review_queue', {}).items():
                for item in items:
                    if item['id'] == item_id:
                        # Update item with review
                        item['status'] = 'reviewed'
                        item['review'] = {
                            'reviewer': data.get('reviewer', 'Unknown'),
                            'timestamp': datetime.now().isoformat(),
                            'decision': data.get('decision'),
                            'corrections': data.get('corrections', {}),
                            'comments': data.get('comments')
                        }
                        
                        # Apply corrections to actual data
                        if data.get('decision') == 'approve' and data.get('corrections'):
                            apply_review_corrections(step, item, data.get('corrections', {}))
                        
                        # Add to history
                        pipeline_state['review_history'].append({
                            'item_id': item_id,
                            'step': step,
                            'timestamp': datetime.now().isoformat(),
                            **item['review']
                        })
                        
                        add_log(f"Review submitted for {item['type']} by {data.get('reviewer')}", 'info')
                        
                        # Emit update via WebSocket
                        socketio.emit('review_update', {
                            'item_id': item_id,
                            'status': 'reviewed',
                            'reviewer': data.get('reviewer')
                        })
                        
                        return jsonify({
                            'status': 'success',
                            'item_id': item_id,
                            'remaining': count_pending_reviews()
                        })
        
        return jsonify({'error': 'Item not found'}), 404
        
    except Exception as e:
        logger.error(f"Review submission error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/review-summary', methods=['GET'])
def get_review_summary():
    """Get overall review status across all steps."""
    try:
        summary = {
            'total_items': 0,
            'reviewed': 0,
            'pending': 0,
            'by_step': {},
            'by_type': {},
            'recent_reviews': [],
            'quality_metrics': calculate_quality_metrics()
        }
        
        with state_lock:
            for step, items in pipeline_state.get('review_queue', {}).items():
                step_summary = {
                    'total': len(items),
                    'reviewed': len([i for i in items if i['status'] == 'reviewed']),
                    'pending': len([i for i in items if i['status'] == 'pending'])
                }
                summary['by_step'][step] = step_summary
                summary['total_items'] += step_summary['total']
                summary['reviewed'] += step_summary['reviewed']
                summary['pending'] += step_summary['pending']
                
                # Count by type
                for item in items:
                    item_type = item['type']
                    if item_type not in summary['by_type']:
                        summary['by_type'][item_type] = {'total': 0, 'reviewed': 0}
                    summary['by_type'][item_type]['total'] += 1
                    if item['status'] == 'reviewed':
                        summary['by_type'][item_type]['reviewed'] += 1
            
            # Get recent reviews
            summary['recent_reviews'] = sorted(
                pipeline_state.get('review_history', []),
                key=lambda x: x['timestamp'],
                reverse=True
            )[:10]
        
        return jsonify(summary)
        
    except Exception as e:
        logger.error(f"Review summary error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/batch-review', methods=['POST'])
def batch_review():
    """Review multiple items at once."""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        item_ids = data.get('item_ids', [])
        reviewer = data.get('reviewer', 'Unknown')
        decision = data.get('decision', 'approve')
        
        reviewed_count = 0
        
        with state_lock:
            for step, items in pipeline_state.get('review_queue', {}).items():
                for item in items:
                    if item['id'] in item_ids and item['status'] == 'pending':
                        item['status'] = 'reviewed'
                        item['review'] = {
                            'reviewer': reviewer,
                            'timestamp': datetime.now().isoformat(),
                            'decision': decision,
                            'batch': True,
                            'comments': f"Batch {decision}"
                        }
                        
                        pipeline_state['review_history'].append({
                            'item_id': item['id'],
                            'step': step,
                            'timestamp': datetime.now().isoformat(),
                            **item['review']
                        })
                        
                        reviewed_count += 1
        
        add_log(f"Batch review: {reviewed_count} items {decision}ed by {reviewer}", 'info')
        
        socketio.emit('batch_review_complete', {
            'count': reviewed_count,
            'decision': decision,
            'reviewer': reviewer
        })
        
        return jsonify({
            'status': 'success',
            'reviewed_count': reviewed_count,
            'remaining': count_pending_reviews()
        })
        
    except Exception as e:
        logger.error(f"Batch review error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/export/review-report', methods=['GET'])
def export_review_report():
    """Generate a comprehensive review report."""
    try:
        with state_lock:
            report = {
                'generated_at': datetime.now().isoformat(),
                'session_id': pipeline_state['current_session'],
                'review_summary': {
                    'total_items_reviewed': len(pipeline_state['review_history']),
                    'reviewers': list(set(r['reviewer'] for r in pipeline_state['review_history'])),
                    'pending_reviews': count_pending_reviews()
                },
                'quality_metrics': calculate_quality_metrics(),
                'review_details': pipeline_state['review_history']
            }
        
        # Create temporary file
        temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False)
        json.dump(report, temp_file, indent=2)
        temp_file.close()
        
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name=f'review_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.json',
            mimetype='application/json'
        )
        
    except Exception as e:
        logger.error(f"Review report error: {e}")
        return jsonify({'error': str(e)}), 500

# ==================== ORIGINAL ENDPOINTS WITH REVIEW INTEGRATION ====================

def validate_json_structure(data, step):
    """Validate JSON structure for each step."""
    errors = []
    warnings = []
    
    try:
        if step == 2:  # DFD
            if 'dfd' not in data:
                errors.append("Missing 'dfd' key in output")
            else:
                dfd = data['dfd']
                required_fields = ['project_name', 'external_entities', 'processes', 'assets', 'data_flows']
                for field in required_fields:
                    if field not in dfd:
                        errors.append(f"Missing required field: {field}")
                    elif not dfd[field]:
                        warnings.append(f"Empty field: {field}")
                
                # Validate data flows
                if 'data_flows' in dfd and isinstance(dfd['data_flows'], list):
                    for i, flow in enumerate(dfd['data_flows']):
                        if not isinstance(flow, dict):
                            errors.append(f"Data flow {i} is not a dictionary")
                        else:
                            for req in ['source', 'destination']:
                                if req not in flow:
                                    errors.append(f"Data flow {i} missing '{req}'")
        
        elif step == 3 or step == 4:  # Threats
            if 'threats' not in data:
                errors.append("Missing 'threats' key in output")
            else:
                threats = data['threats']
                if not isinstance(threats, list):
                    errors.append("'threats' must be a list")
                else:
                    for i, threat in enumerate(threats):
                        if not isinstance(threat, dict):
                            errors.append(f"Threat {i} is not a dictionary")
                        else:
                            required = ['component_name', 'stride_category', 'threat_description', 
                                      'mitigation_suggestion', 'impact', 'likelihood']
                            for field in required:
                                if field not in threat:
                                    errors.append(f"Threat {i} missing '{field}'")
        
        elif step == 5:  # Attack Paths
            if 'attack_paths' not in data:
                errors.append("Missing 'attack_paths' key in output")
            else:
                paths = data['attack_paths']
                if not isinstance(paths, list):
                    errors.append("'attack_paths' must be a list")
    
    except Exception as e:
        errors.append(f"Validation error: {str(e)}")
    
    return {
        'valid': len(errors) == 0,
        'errors': errors,
        'warnings': warnings
    }

# Add request logging and error handling
@app.before_request
def log_request_info():
    logger.info(f"Request: {request.method} {request.url}")
    if request.method == 'POST':
        logger.info(f"Content-Type: {request.content_type}")
        if request.is_json:
            logger.info(f"JSON body: {request.get_json()}")
        else:
            logger.info(f"Form data keys: {list(request.form.keys())}")
            logger.info(f"Files: {list(request.files.keys())}")

@app.errorhandler(400)
def handle_bad_request(e):
    logger.error(f"400 Bad Request: {e}")
    return jsonify({'error': 'Bad Request', 'message': str(e)}), 400

@app.errorhandler(500)
def handle_internal_error(e):
    logger.error(f"500 Internal Server Error: {e}")
    logger.error(f"Traceback: {traceback.format_exc()}")
    return jsonify({'error': 'Internal Server Error', 'message': str(e)}), 500

@app.route('/api/test', methods=['GET', 'POST'])
def test_endpoint():
    """Simple test endpoint for debugging."""
    return jsonify({
        'status': 'success',
        'method': request.method,
        'timestamp': datetime.now().isoformat(),
        'message': 'API is working',
        'config': {
            'llm_provider': runtime_config['llm_provider'],
            'has_api_key': bool(runtime_config['scw_secret_key']),
            'output_dir': OUTPUT_FOLDER,
            'input_dir': INPUT_FOLDER
        }
    })

@app.route('/api/configuration', methods=['GET', 'POST'])
def handle_configuration():
    """Get or update runtime configuration."""
    global runtime_config
    
    if request.method == 'GET':
        return jsonify(runtime_config)
    
    elif request.method == 'POST':
        try:
            updates = request.get_json()
            if not updates:
                return jsonify({'error': 'No JSON data provided'}), 400
            
            # Update runtime configuration
            for key, value in updates.items():
                if key in runtime_config:
                    runtime_config[key] = value
                    add_log(f"Configuration updated: {key} = {value}", 'info')
            
            return jsonify({
                'status': 'updated',
                'configuration': runtime_config
            })
        except Exception as e:
            logger.error(f"Configuration update error: {e}")
            return jsonify({'error': str(e)}), 500

@app.route('/api/upload', methods=['POST'])
def upload_document():
    """Handle document upload and text extraction."""
    try:
        logger.info(f"Upload request received")
        
        if 'document' not in request.files:
            logger.error("No 'document' field in request files")
            return jsonify({'error': 'No file part in request'}), 400
        
        file = request.files['document']
        logger.info(f"File received: {file.filename}")
        
        if file.filename == '':
            logger.error("Empty filename")
            return jsonify({'error': 'No file selected'}), 400
        
        if file and allowed_file(file.filename):
            # Save uploaded file
            filename = secure_filename(file.filename)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"{timestamp}_{filename}"
            file_path = os.path.join(UPLOAD_FOLDER, filename)
            file.save(file_path)
            
            add_log(f"File uploaded: {filename}", 'success')
            
            # Extract text content
            text_content, error = extract_text_from_file(file_path)
            
            if error:
                add_log(f"Text extraction failed: {error}", 'error')
                return jsonify({'error': f'Failed to extract text: {error}'}), 500
            
            if not text_content or len(text_content.strip()) < 10:
                add_log("Extracted text is too short or empty", 'error')
                return jsonify({'error': 'Extracted text is empty or too short'}), 400
            
            # Save extracted text for the pipeline
            text_file_path = os.path.join(INPUT_FOLDER, f"{timestamp}_extracted.txt")
            with open(text_file_path, 'w', encoding='utf-8') as f:
                f.write(text_content)
            
            add_log(f"Text extracted: {len(text_content)} characters", 'success')
            
            # Create session and save to pipeline state
            session_id = timestamp
            upload_data = {
                'status': 'success',
                'session_id': session_id,
                'filename': filename,
                'file_path': file_path,
                'text_file_path': text_file_path,
                'text_preview': text_content[:500] + '...' if len(text_content) > 500 else text_content,
                'text_length': len(text_content),
                'count': 1
            }
            
            with state_lock:
                pipeline_state['current_session'] = session_id
                pipeline_state['step_outputs'][1] = upload_data
                if 'logs' not in pipeline_state:
                    pipeline_state['logs'] = []
                if 'validations' not in pipeline_state:
                    pipeline_state['validations'] = {}
            
            add_log(f"Step 1 completed successfully: {filename}", 'success')
            
            return jsonify(upload_data)
        
        logger.error(f"File type not allowed: {file.filename}")
        return jsonify({'error': 'File type not allowed. Please use TXT, PDF, or DOCX files.'}), 400
        
    except Exception as e:
        logger.error(f"Upload error: {str(e)}\n{traceback.format_exc()}")
        add_log(f"Upload error: {str(e)}", 'error')
        return jsonify({'error': f'Upload failed: {str(e)}'}), 500

@app.route('/api/run-step', methods=['POST'])
def run_step():
    """Run a specific pipeline step with review generation."""
    try:
        logger.info("RUN-STEP REQUEST RECEIVED")
        
        if not request.is_json:
            return jsonify({'error': 'Request must be JSON'}), 400
        
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No JSON data provided'}), 400
        
        if 'step' not in data:
            return jsonify({'error': 'Missing step parameter'}), 400
        
        step = data['step']
        input_data = data.get('input', {})
        
        logger.info(f"Running step {step}")
        
        # Validate step number
        if not isinstance(step, int) or step < 1 or step > 5:
            return jsonify({'error': 'Invalid step number. Must be 1-5.'}), 400
        
        # Check prerequisites
        if step > 1:
            with state_lock:
                step_outputs = pipeline_state.get('step_outputs', {})
                prev_step_output = step_outputs.get(step - 1)
                
                if not prev_step_output:
                    # Try to recover from disk
                    if step == 3:
                        dfd_file = os.path.join(OUTPUT_FOLDER, 'dfd_components.json')
                        if os.path.exists(dfd_file):
                            try:
                                with open(dfd_file, 'r') as f:
                                    pipeline_state['step_outputs'][2] = json.load(f)
                                logger.info("Step 2 data recovered from disk")
                            except Exception as e:
                                logger.error(f"Failed to recover step 2 data: {e}")
                                return jsonify({'error': f'Step {step - 1} must be completed first'}), 400
                        else:
                            return jsonify({'error': f'Step {step - 1} must be completed first'}), 400
        
        add_log(f"Starting step {step}", 'info')
        
        # Setup environment
        env = os.environ.copy()
        env_updates = {
            'INPUT_DIR': INPUT_FOLDER,
            'OUTPUT_DIR': OUTPUT_FOLDER,
            'LOG_LEVEL': 'INFO',
            'PIPELINE_TIMEOUT': str(runtime_config['timeout']),
            'LLM_PROVIDER': runtime_config['llm_provider'],
            'LLM_MODEL': runtime_config['llm_model'],
            'LOCAL_LLM_ENDPOINT': runtime_config['local_llm_endpoint'],
            'CUSTOM_SYSTEM_PROMPT': runtime_config['custom_system_prompt'],
            'MITRE_ENABLED': str(runtime_config['mitre_enabled']).lower(),
            'MITRE_VERSION': runtime_config['mitre_version'],
            'SCW_API_URL': runtime_config['scw_api_url']
        }
        
        env.update(env_updates)
        
        if runtime_config['scw_secret_key']:
            env['SCW_SECRET_KEY'] = runtime_config['scw_secret_key']
            env['SCW_API_KEY'] = runtime_config['scw_secret_key']
        
        result = None
        output_file = None
        script_name = None
        
        if step == 1:
            # Document already uploaded
            if 1 not in pipeline_state.get('step_outputs', {}):
                return jsonify({'error': 'Please upload a document first'}), 400
            return jsonify(pipeline_state['step_outputs'][1])
            
        elif step == 2:
            script_name = 'info_to_dfds.py'
            env['DFD_OUTPUT_PATH'] = os.path.join(OUTPUT_FOLDER, 'dfd_components.json')
            output_file = env['DFD_OUTPUT_PATH']
            
        elif step == 3:
            script_name = 'dfd_to_threats.py'
            env['DFD_INPUT_PATH'] = os.path.join(OUTPUT_FOLDER, 'dfd_components.json')
            env['THREATS_OUTPUT_PATH'] = os.path.join(OUTPUT_FOLDER, 'identified_threats.json')
            output_file = env['THREATS_OUTPUT_PATH']
            
        elif step == 4:
            script_name = 'improve_threat_quality.py'
            env.update({
                'DFD_INPUT_PATH': os.path.join(OUTPUT_FOLDER, 'dfd_components.json'),
                'THREATS_INPUT_PATH': os.path.join(OUTPUT_FOLDER, 'identified_threats.json'),
                'REFINED_THREATS_OUTPUT_PATH': os.path.join(OUTPUT_FOLDER, 'refined_threats.json'),
                'SIMILARITY_THRESHOLD': '0.80',
                'CVE_RELEVANCE_YEARS': '5'
            })
            output_file = env['REFINED_THREATS_OUTPUT_PATH']
            
        elif step == 5:
            script_name = 'attack_path_analyzer.py'
            env.update({
                'REFINED_THREATS_PATH': os.path.join(OUTPUT_FOLDER, 'refined_threats.json'),
                'DFD_PATH': os.path.join(OUTPUT_FOLDER, 'dfd_components.json'),
                'ATTACK_PATHS_OUTPUT': os.path.join(OUTPUT_FOLDER, 'attack_paths.json'),
                'MAX_PATH_LENGTH': '5',
                'MAX_PATHS_TO_ANALYZE': '20',
                'ENABLE_VECTOR_STORE': 'false'
            })
            output_file = env['ATTACK_PATHS_OUTPUT']
        
        # Check script exists
        if script_name and not os.path.exists(script_name):
            return jsonify({'error': f'Script not found: {script_name}'}), 500
        
        # Run the script
        if script_name:
            result = subprocess.run(
                [sys.executable, script_name],
                capture_output=True,
                text=True,
                env=env,
                timeout=int(runtime_config['timeout']),
                cwd=os.getcwd()
            )
            
            if result.returncode != 0:
                error_msg = result.stderr or result.stdout or f'Script {script_name} failed'
                add_log(f"Step {step} failed: {error_msg}", 'error')
                return jsonify({'error': f'Script execution failed: {error_msg}'}), 500
        
        # Load and validate output
        if output_file and os.path.exists(output_file):
            try:
                with open(output_file, 'r', encoding='utf-8') as f:
                    step_data = json.load(f)
                
                # Validate structure
                validation = validate_json_structure(step_data, step)
                
                if not validation['valid']:
                    add_log(f"Step {step} validation errors: {', '.join(validation['errors'])}", 'warning')
                
                with state_lock:
                    pipeline_state['validations'][step] = validation
                    pipeline_state['step_outputs'][step] = step_data
                    
                    # Generate review items for this step
                    review_items = generate_review_items(step, step_data)
                    items_needing_review = [
                        item for item in review_items 
                        if item['confidence'] < 0.8 or 'missing_fields' in item or 'attributes_needed' in item
                    ]
                    
                    if items_needing_review:
                        pipeline_state['review_queue'][step] = items_needing_review
                        
                        # Notify frontend about pending reviews
                        socketio.emit('reviews_available', {
                            'step': step,
                            'count': len(items_needing_review)
                        })
                
                # Calculate count for response
                count = 0
                if step == 2 and 'dfd' in step_data:
                    dfd = step_data['dfd']
                    count = (len(dfd.get('processes', [])) + 
                            len(dfd.get('assets', [])) +
                            len(dfd.get('external_entities', [])))
                elif step in [3, 4] and 'threats' in step_data:
                    count = len(step_data['threats'])
                elif step == 5 and 'attack_paths' in step_data:
                    count = len(step_data['attack_paths'])
                
                add_log(f"Step {step} completed: {count} items found", 'success')
                
                response_data = {
                    **step_data,
                    'count': count,
                    'validation': validation,
                    'step': step,
                    'timestamp': datetime.now().isoformat(),
                    'review_needed': len(items_needing_review) if 'items_needing_review' in locals() else 0
                }
                
                return jsonify(response_data)
                
            except json.JSONDecodeError as e:
                add_log(f"Step {step} output is not valid JSON: {e}", 'error')
                return jsonify({'error': f'Script output is not valid JSON: {e}'}), 500
                
            except Exception as e:
                add_log(f"Step {step} failed to process output: {e}", 'error')
                return jsonify({'error': f'Failed to process output: {e}'}), 500
        else:
            add_log(f"Step {step} output file not created: {output_file}", 'error')
            return jsonify({'error': f'Output file not created: {output_file}'}), 500
            
    except subprocess.TimeoutExpired:
        add_log(f"Step {step} timed out after {runtime_config['timeout']} seconds", 'error')
        return jsonify({'error': f'Script execution timed out'}), 500
    except Exception as e:
        logger.error(f"Step execution error: {str(e)}\n{traceback.format_exc()}")
        add_log(f"Step execution error: {str(e)}", 'error')
        return jsonify({'error': f'Step execution failed: {str(e)}'}), 500

@app.route('/api/save-step', methods=['POST'])
def save_step():
    """Save edited data for a specific step."""
    try:
        if not request.is_json:
            return jsonify({'error': 'Request must be JSON'}), 400
        
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No JSON data provided'}), 400
        
        step = data.get('step')
        step_data = data.get('data')
        
        if not step or not step_data:
            return jsonify({'error': 'Missing step or data parameter'}), 400
        
        # Validate the new data
        validation = validate_json_structure(step_data, step)
        if not validation['valid']:
            return jsonify({
                'error': 'Invalid data structure',
                'validation': validation
            }), 400
        
        # Save to file
        save_step_data(step, step_data)
        
        # Update state
        with state_lock:
            pipeline_state['step_outputs'][step] = step_data
            pipeline_state['validations'][step] = validation
            
            # Regenerate review items
            review_items = generate_review_items(step, step_data)
            if review_items:
                pipeline_state['review_queue'][step] = review_items
        
        add_log(f"Saved changes to step {step}", 'success')
        return jsonify({
            'status': 'saved',
            'validation': validation
        })
        
    except Exception as e:
        logger.error(f"Save error: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/export/<step>', methods=['GET'])
def export_data(step):
    """Export data for a specific step."""
    try:
        files = {
            '2': 'dfd_components.json',
            '3': 'identified_threats.json',
            '4': 'refined_threats.json',
            '5': 'attack_paths.json',
            'all': 'complete_analysis.json'
        }
        
        if step == 'all':
            # Combine all data
            complete_data = {
                'export_date': datetime.now().isoformat(),
                'session_id': pipeline_state['current_session'],
                'validations': pipeline_state.get('validations', {}),
                'review_summary': {
                    'total_reviews': len(pipeline_state.get('review_history', [])),
                    'pending_reviews': count_pending_reviews()
                },
                'steps': {}
            }
            
            for step_num in range(1, 6):
                if step_num in pipeline_state.get('step_outputs', {}):
                    complete_data['steps'][step_num] = pipeline_state['step_outputs'][step_num]
            
            # Create temporary file
            temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False)
            json.dump(complete_data, temp_file, indent=2)
            temp_file.close()
            
            return send_file(
                temp_file.name, 
                as_attachment=True, 
                download_name=f'threat_model_{datetime.now().strftime("%Y%m%d_%H%M%S")}.json',
                mimetype='application/json'
            )
        
        elif step in files:
            file_path = os.path.join(OUTPUT_FOLDER, files[step])
            if os.path.exists(file_path):
                return send_file(
                    file_path, 
                    as_attachment=True,
                    download_name=f'{files[step].replace(".json", "")}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.json',
                    mimetype='application/json'
                )
            else:
                return jsonify({'error': 'File not found'}), 404
        
        return jsonify({'error': 'Invalid step'}), 400
        
    except Exception as e:
        logger.error(f"Export error: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/logs', methods=['GET'])
def get_logs():
    """Get pipeline execution logs."""
    try:
        with state_lock:
            limit = request.args.get('limit', 100, type=int)
            logs = pipeline_state['logs'][-limit:]
        
        return jsonify({'logs': logs})
    except Exception as e:
        logger.error(f"Error getting logs: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/status', methods=['GET'])
def get_status():
    """Get current pipeline status."""
    try:
        with state_lock:
            completed_steps = []
            for step in range(1, 6):
                if step in pipeline_state.get('step_outputs', {}):
                    completed_steps.append(step)
            
            status = {
                'session_id': pipeline_state['current_session'],
                'completed_steps': completed_steps,
                'validations': pipeline_state.get('validations', {}),
                'pending_reviews': count_pending_reviews(),
                'last_log': pipeline_state['logs'][-1] if pipeline_state['logs'] else None,
                'timestamp': datetime.now().isoformat()
            }
        
        return jsonify(status)
    except Exception as e:
        logger.error(f"Error getting status: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/step-progress/<int:step>', methods=['GET'])
def get_step_progress(step):
    """Get real-time progress for a running step."""
    try:
        progress_file = os.path.join(OUTPUT_FOLDER, f'step_{step}_progress.json')
        
        if os.path.exists(progress_file):
            with open(progress_file, 'r') as f:
                progress_data = json.load(f)
            return jsonify(progress_data)
        else:
            return jsonify({
                'status': 'unknown',
                'progress': 0,
                'message': 'No progress information available'
            })
    except Exception as e:
        logger.error(f"Progress check error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint."""
    try:
        return jsonify({
            'status': 'healthy',
            'timestamp': datetime.now().isoformat(),
            'session': pipeline_state['current_session'],
            'has_api_key': bool(runtime_config['scw_secret_key']),
            'scripts_available': all(os.path.exists(s) for s in scripts),
            'review_system': 'enabled',
            'pending_reviews': count_pending_reviews(),
            'directories': {
                'upload': os.path.exists(UPLOAD_FOLDER),
                'output': os.path.exists(OUTPUT_FOLDER),
                'input': os.path.exists(INPUT_FOLDER)
            },
            'dependencies': {
                'pdf_support': PDF_AVAILABLE,
                'docx_support': DOCX_AVAILABLE
            },
            'config': {
                'llm_provider': runtime_config['llm_provider'],
                'llm_model': runtime_config['llm_model'],
                'timeout': runtime_config['timeout']
            }
        })
    except Exception as e:
        logger.error(f"Health check error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/debug-state', methods=['GET'])
def debug_state():
    """Debug endpoint to check pipeline state."""
    try:
        with state_lock:
            debug_info = {
                'current_session': pipeline_state.get('current_session'),
                'completed_steps': list(pipeline_state.get('step_outputs', {}).keys()),
                'review_queue_summary': {
                    step: len(items) for step, items in pipeline_state.get('review_queue', {}).items()
                },
                'total_reviews': len(pipeline_state.get('review_history', [])),
                'files_on_disk': {
                    'dfd_components.json': os.path.exists(os.path.join(OUTPUT_FOLDER, 'dfd_components.json')),
                    'identified_threats.json': os.path.exists(os.path.join(OUTPUT_FOLDER, 'identified_threats.json')),
                    'refined_threats.json': os.path.exists(os.path.join(OUTPUT_FOLDER, 'refined_threats.json')),
                    'attack_paths.json': os.path.exists(os.path.join(OUTPUT_FOLDER, 'attack_paths.json'))
                },
                'timestamp': datetime.now().isoformat()
            }
        
        return jsonify(debug_info)
    except Exception as e:
        logger.error(f"Debug state error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/debug-load-files', methods=['GET', 'POST'])
def debug_load_files():
    """Debug endpoint to manually load existing output files into pipeline state."""
    try:
        files_loaded = {}
        
        # Define file mappings
        file_mappings = {
            2: ('dfd_components.json', 'DFD'),
            3: ('identified_threats.json', 'Threats'),
            4: ('refined_threats.json', 'Refined Threats'),
            5: ('attack_paths.json', 'Attack Paths')
        }
        
        with state_lock:
            # Ensure we have a session
            if not pipeline_state.get('current_session'):
                pipeline_state['current_session'] = datetime.now().strftime('%Y%m%d_%H%M%S')
            
            for step, (filename, description) in file_mappings.items():
                filepath = os.path.join(OUTPUT_FOLDER, filename)
                if os.path.exists(filepath):
                    try:
                        with open(filepath, 'r') as f:
                            data = json.load(f)
                        
                        pipeline_state['step_outputs'][step] = data
                        
                        # Generate review items for loaded data
                        review_items = generate_review_items(step, data)
                        if review_items:
                            pipeline_state['review_queue'][step] = review_items
                        
                        files_loaded[step] = {
                            'status': 'loaded',
                            'file': filename,
                            'description': description,
                            'review_items': len(review_items) if review_items else 0
                        }
                        
                        # Add basic upload data for step 1 if missing
                        if 1 not in pipeline_state['step_outputs'] and step == 2:
                            pipeline_state['step_outputs'][1] = {
                                'status': 'success',
                                'session_id': pipeline_state['current_session'],
                                'filename': 'debug_recovery.txt',
                                'text_length': 1000,
                                'count': 1
                            }
                            files_loaded[1] = {
                                'status': 'simulated',
                                'description': 'Upload step simulated for recovery'
                            }
                        
                        logger.info(f"✅ Loaded {filename} into step {step}")
                        
                    except Exception as e:
                        files_loaded[step] = {
                            'status': 'error',
                            'file': filename,
                            'error': str(e)
                        }
                        logger.error(f"❌ Failed to load {filename}: {e}")
                else:
                    files_loaded[step] = {
                        'status': 'not_found',
                        'file': filename
                    }
        
        return jsonify({
            'status': 'debug_load_complete',
            'files_loaded': files_loaded,
            'current_session': pipeline_state.get('current_session'),
            'completed_steps': list(pipeline_state.get('step_outputs', {}).keys()),
            'pending_reviews': count_pending_reviews()
        })
        
    except Exception as e:
        logger.error(f"Debug load files error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/reset', methods=['POST'])
def reset_pipeline():
    """Reset the pipeline state."""
    global pipeline_state
    
    try:
        with state_lock:
            # Save current session before reset if requested
            if pipeline_state['current_session'] and request.is_json:
                request_data = request.get_json() or {}
                if request_data.get('save_session', False):
                    session_backup = {
                        'session_id': pipeline_state['current_session'],
                        'timestamp': datetime.now().isoformat(),
                        'step_outputs': pipeline_state.get('step_outputs', {}),
                        'validations': pipeline_state.get('validations', {}),
                        'review_history': pipeline_state.get('review_history', [])
                    }
                    backup_file = os.path.join(OUTPUT_FOLDER, f'session_backup_{pipeline_state["current_session"]}.json')
                    with open(backup_file, 'w') as f:
                        json.dump(session_backup, f, indent=2)
                    add_log(f"Session backed up to {os.path.basename(backup_file)}", 'info')
            
            # Reset state
            pipeline_state = {
                'current_session': None,
                'logs': [],
                'step_outputs': {},
                'validations': {},
                'review_queue': {},
                'review_history': [],
                'quality_metrics': {}
            }
        
        # Optionally clean output directory
        if request.is_json:
            request_data = request.get_json() or {}
            if request_data.get('clean_output', False):
                for file in os.listdir(OUTPUT_FOLDER):
                    if file.endswith('.json') and not file.startswith('session_backup'):
                        try:
                            os.remove(os.path.join(OUTPUT_FOLDER, file))
                        except Exception as e:
                            logger.warning(f"Failed to remove {file}: {e}")
        
        add_log("Pipeline reset", 'info')
        return jsonify({'status': 'reset', 'timestamp': datetime.now().isoformat()})
        
    except Exception as e:
        logger.error(f"Reset error: {e}")
        return jsonify({'error': str(e)}), 500

# Add a catch-all error handler
@app.errorhandler(Exception)
def handle_exception(e):
    logger.error(f"Unhandled exception: {e}")
    logger.error(f"Traceback: {traceback.format_exc()}")
    return jsonify({
        'error': 'Internal server error',
        'message': str(e),
        'timestamp': datetime.now().isoformat()
    }), 500

# WebSocket event handlers
@socketio.on('connect')
def handle_connect():
    logger.info('Client connected')
    emit('connected', {'data': 'Connected to review system'})

@socketio.on('disconnect')
def handle_disconnect():
    logger.info('Client disconnected')

if __name__ == '__main__':
    logger.info("Starting Enhanced Threat Modeling Pipeline Backend...")
    logger.info(f"Review system: ENABLED")
    logger.info(f"Upload folder: {UPLOAD_FOLDER}")
    logger.info(f"Output folder: {OUTPUT_FOLDER}")
    logger.info(f"Input folder: {INPUT_FOLDER}")
    
    # Final check for API key
    if not runtime_config['scw_secret_key']:
        logger.warning("⚠️  No API key found! The LLM calls will fail.")
        logger.warning("Please create a .env file with: SCW_SECRET_KEY=your_key_here")
    
    # Create required directories
    for folder in [UPLOAD_FOLDER, OUTPUT_FOLDER, INPUT_FOLDER]:
        try:
            os.makedirs(folder, exist_ok=True)
            logger.info(f"✓ Directory ready: {folder}")
        except Exception as e:
            logger.error(f"✗ Failed to create directory {folder}: {e}")
    
    # Test file permissions
    try:
        test_file = os.path.join(OUTPUT_FOLDER, 'test_permissions.txt')
        with open(test_file, 'w') as f:
            f.write('test')
        os.remove(test_file)
        logger.info("✓ File system permissions OK")
    except Exception as e:
        logger.error(f"✗ File system permission error: {e}")
    
    logger.info("Backend ready with review system!")
    
    # Run with SocketIO
    socketio.run(app, debug=True, port=5000, host='0.0.0.0')