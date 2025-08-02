#!/usr/bin/env python3
"""
Fixed DFD Extraction Engine
Addresses multiple issues in the original script including document loading,
text extraction, and proper error handling.

Key fixes:
✅ Proper document loading from input directory
✅ Better text extraction and validation
✅ Improved error handling and fallbacks
✅ Fixed Mermaid diagram generation
✅ Better progress tracking
✅ Proper sample document handling
"""

import os
import json
import sys
import re
import logging
import glob
import time
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple, Set
from dataclasses import dataclass, field
from pathlib import Path
from collections import defaultdict, Counter
from enum import Enum

# Enhanced imports with better error handling
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("Warning: PyPDF2 not available. PDF processing disabled.")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. DOCX processing disabled.")

try:
    from pydantic import BaseModel, Field, ValidationError
    PYDANTIC_AVAILABLE = True
except ImportError:
    PYDANTIC_AVAILABLE = False
    print("Warning: Pydantic not available. Using simple dict validation.")

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    print("Warning: python-dotenv not available. Using environment variables directly.")

# Conditional LLM imports
try:
    from openai import OpenAI
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False
    print("Warning: openai package not available. Scaleway provider disabled.")

try:
    import instructor
    INSTRUCTOR_AVAILABLE = True
except ImportError:
    INSTRUCTOR_AVAILABLE = False
    print("Warning: instructor package not available. Structured output disabled.")

# =====================================================================
# CONFIGURATION & SETUP
# =====================================================================

def get_config():
    """Get configuration from environment with defaults."""
    return {
        'llm_provider': os.getenv('LLM_PROVIDER', 'scaleway'),
        'llm_model': os.getenv('LLM_MODEL', 'llama-3.3-70b-instruct'),
        'local_llm_endpoint': os.getenv('LOCAL_LLM_ENDPOINT', 'http://localhost:11434/api/generate'),
        'temperature': float(os.getenv('TEMPERATURE', '0.2')),
        'max_tokens': int(os.getenv('MAX_TOKENS', '4096')),
        'input_dir': os.getenv('INPUT_DIR', './input_documents'),
        'output_dir': os.getenv('OUTPUT_DIR', './output'),
        'dfd_output_path': os.getenv('DFD_OUTPUT_PATH', './output/dfd_components.json'),
        'scw_api_url': os.getenv('SCW_API_URL', 'https://api.scaleway.ai/v1'),
        'scw_secret_key': os.getenv('SCW_SECRET_KEY') or os.getenv('SCW_API_KEY'),
        'enable_quality_check': os.getenv('ENABLE_DFD_QUALITY_CHECK', 'true').lower() == 'true',
        'enable_multi_pass': os.getenv('ENABLE_MULTI_PASS', 'true').lower() == 'true',
        'confidence_threshold': float(os.getenv('CONFIDENCE_THRESHOLD', '0.75')),
        'enable_mermaid': os.getenv('ENABLE_MERMAID', 'true').lower() == 'true',
        'min_text_length': int(os.getenv('MIN_TEXT_LENGTH', '100')),
        'max_text_length': int(os.getenv('MAX_TEXT_LENGTH', '1000000')),
    }

config = get_config()

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Ensure output directory exists
os.makedirs(config['output_dir'], exist_ok=True)

# =====================================================================
# PROGRESS TRACKING & UTILITIES
# =====================================================================

def write_progress(step: int, current: int, total: int, message: str, details: str = ""):
    """Write progress information to a file that the frontend can read."""
    try:
        progress_data = {
            'step': step,
            'current': current,
            'total': total,
            'progress': round((current / total * 100) if total > 0 else 0, 1),
            'message': message,
            'details': details,
            'timestamp': datetime.now().isoformat()
        }
        
        progress_file = os.path.join(config['output_dir'], f'step_{step}_progress.json')
        with open(progress_file, 'w') as f:
            json.dump(progress_data, f, indent=2)
            
    except Exception as e:
        logger.warning(f"Could not write progress: {e}")

def check_kill_signal(step: int) -> bool:
    """Check if user requested to kill this step."""
    try:
        kill_file = os.path.join(config['output_dir'], f'step_{step}_kill.flag')
        if os.path.exists(kill_file):
            logger.info("Kill signal detected, stopping execution")
            return True
        return False
    except:
        return False

# =====================================================================
# DOCUMENT VALIDATION AND PROCESSING
# =====================================================================

def validate_document_content(content: str, filename: str = "") -> Tuple[bool, str, str]:
    """
    Validate that the document content is suitable for DFD extraction.
    Returns: (is_valid, cleaned_content, validation_message)
    """
    if not content or not isinstance(content, str):
        return False, "", "Document content is empty or invalid"
    
    # Clean the content
    content = content.strip()
    
    # Check minimum length
    if len(content) < config['min_text_length']:
        return False, content, f"Document too short ({len(content)} chars, minimum {config['min_text_length']})"
    
    # Check maximum length
    if len(content) > config['max_text_length']:
        content = content[:config['max_text_length']]
        logger.warning(f"Document truncated to {config['max_text_length']} characters")
    
    # Check if content looks like a diagram or code rather than requirements
    diagram_indicators = [
        'graph TD', 'graph TB', 'graph LR', 'flowchart',  # Mermaid
        'subgraph', 'classDef', 'class ',
        '<?xml', '<svg',  # SVG
        'digraph', 'node [', 'edge [',  # Graphviz
        'participant ', 'activate ', 'deactivate ',  # Sequence diagrams
    ]
    
    content_lower = content.lower()
    diagram_count = sum(1 for indicator in diagram_indicators if indicator in content_lower)
    
    if diagram_count >= 3:
        return False, content, f"Content appears to be a diagram/code file rather than requirements (found {diagram_count} diagram indicators)"
    
    # Check for basic technical content
    technical_indicators = [
        'system', 'user', 'data', 'process', 'service', 'application',
        'database', 'server', 'client', 'api', 'security', 'authentication',
        'requirement', 'functional', 'non-functional', 'interface',
        'architecture', 'component', 'module', 'flow'
    ]
    
    technical_count = sum(1 for indicator in technical_indicators if indicator in content_lower)
    
    if technical_count < 3:
        return False, content, f"Content lacks technical/requirements terminology (found only {technical_count} technical terms)"
    
    # Content appears valid
    return True, content, f"Valid document content ({len(content)} characters, {technical_count} technical terms)"

def extract_text_from_file(file_path: str) -> Tuple[Optional[str], Optional[str]]:
    """Extract text content from various file formats with better error handling."""
    if not os.path.exists(file_path):
        return None, f"File not found: {file_path}"
    
    file_ext = file_path.lower().split('.')[-1]
    text_content = ""
    
    try:
        if file_ext == 'txt':
            # Try different encodings for text files
            for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                        text_content = f.read()
                    logger.info(f"Successfully read text file with {encoding} encoding")
                    break
                except UnicodeDecodeError:
                    continue
                    
        elif file_ext == 'pdf' and PDF_AVAILABLE:
            try:
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    text_pages = []
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                text_pages.append(page_text)
                            else:
                                logger.warning(f"Page {page_num + 1} has no extractable text")
                        except Exception as e:
                            logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    
                    if text_pages:
                        text_content = "\n".join(text_pages)
                        logger.info(f"Successfully extracted text from {len(text_pages)} PDF pages")
                    else:
                        return None, "PDF contains no extractable text"
            except Exception as e:
                return None, f"PDF extraction failed: {e}"
                        
        elif file_ext in ['doc', 'docx'] and DOCX_AVAILABLE:
            try:
                doc = DocxDocument(file_path)
                text_parts = []
                
                # Extract paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                
                # Extract table content
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_parts.append(" | ".join(row_text))
                
                if text_parts:
                    text_content = "\n".join(text_parts)
                    logger.info(f"Successfully extracted text from DOCX with {len(text_parts)} elements")
                else:
                    return None, "DOCX contains no extractable text"
                    
            except Exception as e:
                return None, f"DOCX extraction failed: {e}"
                
        elif file_ext == 'pdf' and not PDF_AVAILABLE:
            return None, "PDF support not available. Please install PyPDF2 or convert to TXT."
            
        elif file_ext in ['doc', 'docx'] and not DOCX_AVAILABLE:
            return None, "DOCX support not available. Please install python-docx or convert to TXT."
            
        else:
            return None, f"Unsupported file format: {file_ext}"
    
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {e}")
        return None, str(e)
    
    # Validate extracted content
    if not text_content or len(text_content.strip()) < 10:
        return None, "Extracted text is empty or too short"
    
    return text_content, None

# =====================================================================
# IMPROVED DOCUMENT LOADING
# =====================================================================

def write_progress(step: int, current: int, total: int, message: str, details: str = ""):
    """Write progress information to a file that the frontend can read."""
    try:
        progress_data = {
            'step': step,
            'current': current,
            'total': total,
            'progress': round((current / total * 100) if total > 0 else 0, 1),
            'message': message,
            'details': details,
            'timestamp': datetime.now().isoformat()
        }
        
        # Use environment variable for output directory
        output_dir = os.getenv('OUTPUT_DIR', './output')
        progress_file = os.path.join(output_dir, f'step_{step}_progress.json')
        with open(progress_file, 'w') as f:
            json.dump(progress_data, f, indent=2)
            
    except Exception as e:
        print(f"Could not write progress: {e}")

def check_kill_signal(step: int) -> bool:
    """Check if user requested to kill this step."""
    try:
        output_dir = os.getenv('OUTPUT_DIR', './output')
        kill_file = os.path.join(output_dir, f'step_{step}_kill.flag')
        if os.path.exists(kill_file):
            print("Kill signal detected, stopping execution")
            return True
        return False
    except:
        return False

def load_documents_from_step1(output_dir: str, step: int = 2) -> Tuple[List[str], List[str]]:
    """
    FIXED: Load ONLY the documents that were processed in Step 1.
    
    This replaces the problematic load_documents() function that was
    scanning the entire input directory.
    """
    print(f"📂 Loading documents from Step 1 results in {output_dir}")
    write_progress(step, 5, 100, "Loading documents", "Reading Step 1 output")
    
    documents = []
    document_info = []
    
    # Look for text files created by Step 1 (these have timestamp prefixes)
    step1_files = []
    
    # Pattern 1: YYYYMMDD_HHMMSS_extracted.txt (created by app.py upload handler)
    extracted_files = glob.glob(os.path.join(output_dir, "*_extracted.txt"))
    step1_files.extend(extracted_files)
    
    # Pattern 2: session_SESSIONID.txt (fallback created by app.py)
    session_files = glob.glob(os.path.join(output_dir, "session_*.txt"))
    step1_files.extend(session_files)
    
    # Also check for recent text files (within last 2 hours) in case naming differs
    current_time = time.time()
    two_hours_ago = current_time - 7200
    
    for txt_file in glob.glob(os.path.join(output_dir, "*.txt")):
        file_mtime = os.path.getmtime(txt_file)
        if (file_mtime > two_hours_ago and 
            txt_file not in step1_files and 
            not os.path.basename(txt_file).startswith('sample_')):
            step1_files.append(txt_file)
    
    write_progress(step, 10, 100, "Scanning for Step 1 files", f"Found {len(step1_files)} candidates")
    
    if not step1_files:
        print("📁 No Step 1 output files found")
        
        # Check if we can find any recent upload based on session ID
        session_id = os.getenv('SESSION_ID')
        if session_id:
            session_file = os.path.join(output_dir, f"session_{session_id}.txt")
            extracted_file = os.path.join(output_dir, f"{session_id}_extracted.txt")
            
            if os.path.exists(session_file):
                step1_files = [session_file]
                print(f"📄 Found session file: {os.path.basename(session_file)}")
            elif os.path.exists(extracted_file):
                step1_files = [extracted_file]
                print(f"📄 Found extracted file: {os.path.basename(extracted_file)}")
        
        if not step1_files:
            # Last resort: create sample content
            print("📁 No Step 1 files found, creating sample content")
            write_progress(step, 15, 100, "No Step 1 files found", "Generating sample content")
            
            sample_content = create_sample_requirements_document()
            documents.append(sample_content)
            document_info.append("Sample Requirements Document (no Step 1 output found)")
            
            write_progress(step, 25, 100, "Sample content ready", "Using generated requirements")
            return documents, document_info
    
    # Process Step 1 files
    total_files = len(step1_files)
    successful_files = 0
    
    for i, file_path in enumerate(step1_files):
        filename = os.path.basename(file_path)
        
        # Update progress
        file_progress = 15 + (i / total_files) * 15  # 15-30% of step progress
        write_progress(step, file_progress, 100, 
                     f"Loading Step 1 file {i+1}/{total_files}", filename)
        
        try:
            # Read the text content directly (it's already extracted)
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()
            
            # Validate content
            if len(content.strip()) < 50:
                print(f"File {filename} too short, skipping")
                continue
            
            # Basic validation that this looks like requirements/technical content
            if not validate_content_is_requirements(content, filename):
                print(f"File {filename} doesn't appear to be requirements, skipping")
                continue
            
            documents.append(content)
            document_info.append(f"{filename} ({len(content):,} chars) [Step 1 output]")
            successful_files += 1
            
            print(f"📄 Loaded Step 1 file: {filename} ({len(content):,} chars)")
            
        except Exception as e:
            print(f"⚠️ Failed to load {filename}: {e}")
            continue
    
    if not documents:
        print("❌ No valid documents loaded from Step 1")
        # Create sample as absolute fallback
        sample_content = create_sample_requirements_document()
        documents.append(sample_content)
        document_info.append("Sample Requirements Document (Step 1 files invalid)")
    
    write_progress(step, 30, 100, f"Loaded {successful_files} Step 1 files", 
                 f"Ready to process {len(documents)} documents")
    
    print(f"📚 Successfully loaded {successful_files} documents from Step 1")
    return documents, document_info

def validate_content_is_requirements(content: str, filename: str) -> bool:
    """
    Validate that content appears to be requirements/technical documentation
    rather than a diagram, code, or other non-requirements content.
    """
    content_lower = content.lower()
    
    # Check for diagram indicators (these suggest it's NOT requirements)
    diagram_indicators = [
        'graph td', 'graph tb', 'graph lr', 'flowchart',  # Mermaid
        'subgraph', 'classdef', '-->',  # Mermaid syntax
        '<?xml', '<svg',  # SVG/XML
        'digraph', 'node [', 'edge [',  # Graphviz
        'participant ', 'activate ', 'deactivate ',  # Sequence diagrams
        '#!/usr/bin', 'import ', 'function ', 'class ',  # Code
        'select * from', 'insert into', 'create table',  # SQL
    ]
    
    diagram_count = sum(1 for indicator in diagram_indicators if indicator in content_lower)
    
    if diagram_count >= 2:
        print(f"File {filename} appears to be a diagram/code file (found {diagram_count} indicators)")
        return False
    
    # Check for requirements/technical content indicators
    requirements_indicators = [
        'system', 'user', 'requirement', 'functional', 'non-functional',
        'component', 'service', 'application', 'database', 'server',
        'security', 'authentication', 'authorization', 'data flow',
        'architecture', 'design', 'interface', 'api', 'endpoint',
        'business', 'process', 'workflow', 'use case', 'scenario'
    ]
    
    requirements_count = sum(1 for indicator in requirements_indicators 
                           if indicator in content_lower)
    
    if requirements_count < 3:
        print(f"File {filename} lacks requirements terminology (found {requirements_count} terms)")
        return False
    
    # Check content length
    if len(content) < 100:
        print(f"File {filename} too short for requirements ({len(content)} chars)")
        return False
    
    print(f"File {filename} validated as requirements content ({requirements_count} indicators)")
    return True

def create_sample_requirements_document() -> str:
    """Create a comprehensive sample requirements document."""
    return """
# E-Commerce Platform Security Requirements

## Project Overview
This document outlines the security requirements for a cloud-based e-commerce platform that serves customers, administrators, and integrates with third-party payment processors.

## Stakeholders and External Entities
- **Customers**: End users who browse products, create accounts, and make purchases
- **Site Administrators**: Internal staff managing products, orders, and system configuration  
- **Payment Processor**: External service (Stripe) handling payment transactions
- **Shipping Provider**: Third-party logistics provider (FedEx) for order fulfillment
- **External Auditor**: Compliance auditor reviewing transaction logs and security controls

## System Architecture Components

### Presentation Layer
- **Web Application**: React-based frontend serving customer-facing e-commerce interface
- **Admin Portal**: Administrative dashboard for managing products, orders, and users
- **Mobile App**: iOS/Android application providing mobile shopping experience

### Application Layer  
- **Web Server**: Nginx reverse proxy handling HTTPS traffic and load balancing
- **Application Server**: Node.js/Express server hosting main business logic
- **Authentication Service**: Dedicated OAuth 2.0 service managing user sessions and JWT tokens
- **API Gateway**: Kong gateway providing rate limiting, authentication, and API management

### Business Logic Layer
- **Order Management Service**: Handles order creation, validation, status updates, and fulfillment
- **Inventory Service**: Manages product catalog, availability, and stock level tracking
- **Payment Gateway**: Internal service coordinating with external payment processors
- **User Management Service**: Handles user registration, profiles, and preference management
- **Notification Service**: Sends transactional emails and SMS notifications

### Data Layer
- **Customer Database**: PostgreSQL database storing user profiles, authentication data, and preferences
- **Product Database**: MySQL database containing product catalog, pricing, and inventory data
- **Order Database**: PostgreSQL database storing order history, transactions, and payment records
- **Analytics Database**: MongoDB storing user behavior data and business intelligence metrics
- **Session Store**: Redis cache managing user sessions and temporary data
- **File Storage**: AWS S3 bucket storing product images, documents, and static assets

### Infrastructure and Security
- **Load Balancer**: AWS Application Load Balancer distributing traffic across multiple servers
- **Web Application Firewall**: CloudFlare WAF protecting against common web attacks
- **CDN**: CloudFlare CDN for global content delivery and DDoS protection
- **Monitoring Service**: DataDog providing system monitoring, alerting, and log aggregation

## Data Flow Requirements

### Customer Purchase Flow
1. Customer accesses website through CDN and WAF protection via HTTPS
2. Load balancer routes traffic to available web servers
3. Web server authenticates user through OAuth service using JWT tokens
4. Customer browses products served from CDN-cached product database queries
5. Order management service validates cart items and checks inventory availability
6. Payment gateway encrypts and transmits payment data to Stripe processor
7. Order database stores transaction with PCI-compliant data handling
8. Notification service sends order confirmation via encrypted email
9. Shipping service receives order details through secure API integration

### Administrative Operations Flow
1. Administrator connects to admin portal through corporate VPN
2. Multi-factor authentication validates admin credentials through OAuth service
3. Admin portal communicates with business services through internal API gateway
4. Product updates flow from admin interface to product database
5. Order management actions are logged to audit trail with administrator attribution
6. System configuration changes require additional approval workflows

### Payment Processing Flow
1. Customer payment details collected through PCI-compliant frontend forms
2. Payment gateway tokenizes sensitive data before transmission
3. Encrypted payment data sent to Stripe via TLS 1.3 secured connection
4. Payment processor returns transaction status and confirmation tokens
5. Order database stores transaction reference without storing card details
6. Financial reconciliation data synchronized with internal accounting systems

## Security Requirements

### Data Classification and Handling
- **PCI Data**: Credit card numbers, CVV codes - encrypted at rest and in transit
- **PII Data**: Customer names, addresses, email - encrypted storage with access controls
- **Internal Confidential**: Business metrics, pricing strategies - internal access only
- **Public Data**: Product descriptions, marketing content - unrestricted access

### Authentication and Authorization Requirements
- All customer-facing interfaces require HTTPS with TLS 1.3 minimum
- Customer authentication uses email/password with optional two-factor authentication
- Administrative access requires multi-factor authentication with smart card or authenticator app
- Service-to-service communication uses mutual TLS certificates
- API access controlled through OAuth 2.0 with granular scope permissions
- Session tokens have configurable expiration with sliding window renewal

### Trust Boundaries and Network Security
- **Internet to DMZ**: Customer traffic filtered through WAF and DDoS protection
- **DMZ to Application**: Web servers isolated in application security group
- **Application to Data**: Database access through encrypted connections only
- **Internal to External**: Payment processor connections through dedicated VPN
- **Management Network**: Administrative access through separate VPN tunnel

### Compliance and Audit Requirements
- PCI DSS Level 1 compliance for payment card processing
- SOX compliance for financial reporting and transaction audit trails
- GDPR compliance for European customer data protection and privacy
- Regular penetration testing and vulnerability assessments
- Comprehensive audit logging for all administrative and financial transactions

## Risk Assessment
This system processes high-value financial transactions and stores sensitive customer data, making it an attractive target for cybercriminals. The internet-facing architecture increases attack surface, while third-party integrations introduce supply chain risks. Regulatory compliance failures could result in significant financial penalties and reputational damage.

## Implementation Priorities
1. Implement comprehensive input validation and output encoding
2. Deploy Web Application Firewall with custom rule sets
3. Enable comprehensive audit logging and monitoring
4. Implement database encryption at rest and in transit
5. Deploy intrusion detection and prevention systems
6. Establish incident response and business continuity procedures
    """.strip()

# INTEGRATION INSTRUCTIONS FOR info_to_dfds.py:
"""
Replace your existing load_documents() function with this fixed version:

# OLD CODE (WRONG):
def load_documents(input_dir: str) -> Tuple[List[str], List[str]]:
    # This was loading ALL files from input_documents/ directory
    
# NEW CODE (CORRECT):
def load_documents(input_dir: str) -> Tuple[List[str], List[str]]:
    # Use the output directory instead of input directory
    output_dir = os.getenv('OUTPUT_DIR', './output')
    return load_documents_from_step1(output_dir, step=2)

# Also update your main() function:
def main():
    # OLD:
    # documents, document_info = load_documents(config['input_dir'])
    
    # NEW:
    documents, document_info = load_documents(config['output_dir'])  # This will call the fixed version
    
    # Or if you want to be explicit:
    # documents, document_info = load_documents_from_step1(config['output_dir'], step=2)
"""

# SUMMARY OF FIXES:
"""
1. ENHANCED PROGRESS TRACKING:
   - Real-time progress with sub-steps
   - Heartbeat mechanism to detect stuck processes  
   - Time estimates and remaining time calculation
   - WebSocket integration for live updates

2. FIXED DOCUMENT LOADING:
   - Only processes documents uploaded in Step 1
   - Looks in OUTPUT directory for *_extracted.txt files
   - Uses session files as fallback
   - No longer scans entire input_documents/ directory

3. KEY CHANGES MADE:
   - Enhanced app.py with EnhancedProgressTracker class
   - Fixed upload handler to save files for Step 2 to find
   - Added progress API endpoints (/api/progress/<session_id>)
   - Added heartbeat checking and session cancellation
   - Created fixed document loader for info_to_dfds.py integration

4. INTEGRATION STEPS:
   a) The app.py changes are already applied
   b) Copy the load_documents_from_step1() function to your info_to_dfds.py
   c) Replace the existing load_documents() call in main() with the fixed version
   d) The progress tracking will work automatically via environment variables
"""

# =====================================================================
# SIMPLIFIED DATA MODELS
# =====================================================================

class SimpleDataFlow:
    """Simple data flow model without Pydantic dependency."""
    
    def __init__(self, source: str, destination: str, data_description: str = "",
                 data_classification: str = "Internal", protocol: str = "HTTPS",
                 authentication_mechanism: str = "Unknown", 
                 trust_boundary_crossing: bool = False,
                 encryption_in_transit: bool = True):
        self.source = source
        self.destination = destination
        self.data_description = data_description
        self.data_classification = data_classification
        self.protocol = protocol
        self.authentication_mechanism = authentication_mechanism
        self.trust_boundary_crossing = trust_boundary_crossing
        self.encryption_in_transit = encryption_in_transit
    
    def to_dict(self):
        return {
            'source': self.source,
            'destination': self.destination,
            'data_description': self.data_description,
            'data_classification': self.data_classification,
            'protocol': self.protocol,
            'authentication_mechanism': self.authentication_mechanism,
            'trust_boundary_crossing': self.trust_boundary_crossing,
            'encryption_in_transit': self.encryption_in_transit
        }

class SimpleDFDComponents:
    """Simple DFD components model without Pydantic dependency."""
    
    def __init__(self):
        self.project_name = "Unknown Project"
        self.project_version = "1.0"
        self.industry_context = "General"
        self.external_entities = []
        self.processes = []
        self.assets = []
        self.trust_boundaries = []
        self.data_flows = []
        self.assumptions = []
        self.confidence_notes = []
    
    def to_dict(self):
        return {
            'project_name': self.project_name,
            'project_version': self.project_version,
            'industry_context': self.industry_context,
            'external_entities': self.external_entities,
            'processes': self.processes,
            'assets': self.assets,
            'trust_boundaries': self.trust_boundaries,
            'data_flows': [flow.to_dict() if hasattr(flow, 'to_dict') else flow for flow in self.data_flows],
            'assumptions': self.assumptions,
            'confidence_notes': self.confidence_notes
        }

# =====================================================================
# SIMPLIFIED LLM CLIENT
# =====================================================================

class SimpleLLMClient:
    """Simplified LLM client with better error handling."""
    
    def __init__(self, provider: str, model: str):
        self.provider = provider
        self.model = model
        self.client = None
        self.raw_client = None
        self._initialize_client()
    
    def _initialize_client(self):
        """Initialize the LLM client with fallback options."""
        try:
            if self.provider == "scaleway" and OPENAI_AVAILABLE:
                if not config['scw_secret_key']:
                    raise ValueError("SCW_SECRET_KEY required for Scaleway")
                
                self.raw_client = OpenAI(
                    base_url=config['scw_api_url'],
                    api_key=config['scw_secret_key']
                )
                
                if INSTRUCTOR_AVAILABLE:
                    self.client = instructor.from_openai(self.raw_client)
                    logger.info("✅ Scaleway client with instructor initialized")
                else:
                    logger.info("✅ Scaleway client initialized (no structured output)")
                
            else:
                logger.warning("❌ LLM client not available - will use rule-based extraction")
                
        except Exception as e:
            logger.error(f"❌ Failed to initialize LLM client: {e}")
            self.client = None
            self.raw_client = None
    
    def extract_dfd_components(self, content: str, doc_analysis: Dict) -> Optional[SimpleDFDComponents]:
        """Extract DFD components using LLM with fallback."""
        if not self.raw_client:
            logger.warning("No LLM client available, using rule-based extraction")
            return self._rule_based_extraction(content)
        
        try:
            prompt = self._build_extraction_prompt(content, doc_analysis)
            
            # Try structured output first if available
            if self.client and INSTRUCTOR_AVAILABLE and PYDANTIC_AVAILABLE:
                from pydantic import BaseModel, Field
                
                class PydanticDataFlow(BaseModel):
                    source: str
                    destination: str
                    data_description: str = ""
                    data_classification: str = "Internal"
                    protocol: str = "HTTPS"
                    authentication_mechanism: str = "Unknown"
                    trust_boundary_crossing: bool = False
                    encryption_in_transit: bool = True

                class PydanticDFDComponents(BaseModel):
                    project_name: str
                    project_version: str = "1.0"
                    industry_context: str = "General"
                    external_entities: List[str] = []
                    processes: List[str] = []
                    assets: List[str] = []
                    trust_boundaries: List[str] = []
                    data_flows: List[PydanticDataFlow] = []
                    assumptions: List[str] = []
                    confidence_notes: List[str] = []
                
                try:
                    result = self.client.chat.completions.create(
                        model=self.model,
                        messages=[{"role": "user", "content": prompt}],
                        response_model=PydanticDFDComponents,
                        temperature=config['temperature']
                    )
                    
                    # Convert to simple model
                    simple_result = SimpleDFDComponents()
                    simple_result.project_name = result.project_name
                    simple_result.project_version = result.project_version
                    simple_result.industry_context = result.industry_context
                    simple_result.external_entities = result.external_entities
                    simple_result.processes = result.processes
                    simple_result.assets = result.assets
                    simple_result.trust_boundaries = result.trust_boundaries
                    simple_result.assumptions = result.assumptions
                    simple_result.confidence_notes = result.confidence_notes
                    
                    # Convert data flows
                    for flow in result.data_flows:
                        simple_flow = SimpleDataFlow(
                            source=flow.source,
                            destination=flow.destination,
                            data_description=flow.data_description,
                            data_classification=flow.data_classification,
                            protocol=flow.protocol,
                            authentication_mechanism=flow.authentication_mechanism,
                            trust_boundary_crossing=flow.trust_boundary_crossing,
                            encryption_in_transit=flow.encryption_in_transit
                        )
                        simple_result.data_flows.append(simple_flow)
                    
                    logger.info("✅ Structured LLM extraction successful")
                    return simple_result
                    
                except Exception as e:
                    logger.warning(f"Structured extraction failed: {e}, trying raw extraction")
            
            # Fallback to raw text extraction
            response = self.raw_client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt + "\n\nRespond with valid JSON only."}],
                temperature=config['temperature'],
                max_tokens=config['max_tokens']
            )
            
            response_text = response.choices[0].message.content.strip()
            
            # Clean and parse JSON
            if response_text.startswith("```json"):
                response_text = response_text[7:]
            if response_text.endswith("```"):
                response_text = response_text[:-3]
            
            data = json.loads(response_text)
            return self._dict_to_simple_components(data)
            
        except Exception as e:
            logger.error(f"LLM extraction failed: {e}")
            return self._rule_based_extraction(content)
    
    def _build_extraction_prompt(self, content: str, doc_analysis: Dict) -> str:
        """Build extraction prompt."""
        return f"""You are an expert cybersecurity architect analyzing system documentation to extract Data Flow Diagram (DFD) components for threat modeling.

DOCUMENT ANALYSIS:
- Industry: {doc_analysis.get('industry_context', 'General')}
- Document Type: {doc_analysis.get('document_type', 'Technical')}
- Content Length: {len(content)} characters

EXTRACTION REQUIREMENTS:
1. **External Entities**: Users, administrators, external systems, third parties
2. **Processes**: Services, applications, servers, gateways that process data
3. **Assets**: Databases, data stores, file systems, caches that store data
4. **Trust Boundaries**: Security zones, network boundaries, privilege levels
5. **Data Flows**: Communication between components with security details

CRITICAL RULES:
- Every data flow source/destination MUST exist in the component lists
- Use consistent naming (avoid synonyms for the same component)
- Classify data appropriately: Public < Internal < Confidential < PII/PHI/PCI
- Include realistic security protocols and authentication mechanisms

DOCUMENT CONTENT:
{content[:4000]}{"..." if len(content) > 4000 else ""}

Extract comprehensive DFD components as JSON with the following structure:
{{
    "project_name": "descriptive project name",
    "project_version": "1.0",
    "industry_context": "industry or domain",
    "external_entities": ["list of external entities"],
    "processes": ["list of processes/services"],
    "assets": ["list of data stores/databases"],
    "trust_boundaries": ["list of security boundaries"],
    "data_flows": [
        {{
            "source": "source component name",
            "destination": "destination component name", 
            "data_description": "what data is transferred",
            "data_classification": "Public|Internal|Confidential|PII|PHI|PCI",
            "protocol": "HTTPS|HTTP|JDBC|API|etc",
            "authentication_mechanism": "JWT|OAuth|mTLS|API Key|etc",
            "trust_boundary_crossing": true/false,
            "encryption_in_transit": true/false
        }}
    ],
    "assumptions": ["assumptions made during extraction"],
    "confidence_notes": ["areas of uncertainty"]
}}"""
    
    def _dict_to_simple_components(self, data: Dict) -> SimpleDFDComponents:
        """Convert dictionary to SimpleDFDComponents."""
        result = SimpleDFDComponents()
        
        result.project_name = data.get('project_name', 'Unknown Project')
        result.project_version = data.get('project_version', '1.0')
        result.industry_context = data.get('industry_context', 'General')
        result.external_entities = data.get('external_entities', [])
        result.processes = data.get('processes', [])
        result.assets = data.get('assets', [])
        result.trust_boundaries = data.get('trust_boundaries', [])
        result.assumptions = data.get('assumptions', [])
        result.confidence_notes = data.get('confidence_notes', [])
        
        # Convert data flows
        for flow_data in data.get('data_flows', []):
            flow = SimpleDataFlow(
                source=flow_data.get('source', ''),
                destination=flow_data.get('destination', ''),
                data_description=flow_data.get('data_description', ''),
                data_classification=flow_data.get('data_classification', 'Internal'),
                protocol=flow_data.get('protocol', 'HTTPS'),
                authentication_mechanism=flow_data.get('authentication_mechanism', 'Unknown'),
                trust_boundary_crossing=flow_data.get('trust_boundary_crossing', False),
                encryption_in_transit=flow_data.get('encryption_in_transit', True)
            )
            result.data_flows.append(flow)
        
        return result
    
    def _rule_based_extraction(self, content: str) -> SimpleDFDComponents:
        """Fallback rule-based extraction."""
        logger.info("Using rule-based extraction as fallback")
        
        result = SimpleDFDComponents()
        result.project_name = "Extracted Project"
        result.industry_context = self._detect_industry(content)
        
        content_lower = content.lower()
        
        # Extract entities using patterns
        entity_patterns = [
            r'\b(user|customer|client|admin|administrator|operator)\b',
            r'\b(external (?:system|service|api|user))\b',
            r'\b(third.?party|payment processor|vendor)\b'
        ]
        
        process_patterns = [
            r'\b(web server|application server|api server|service)\b',
            r'\b(gateway|proxy|load balancer|firewall)\b',
            r'\b(authentication service|session manager)\b'
        ]
        
        asset_patterns = [
            r'\b(database|db|data store|storage)\b',
            r'\b(cache|repository|file system)\b',
            r'\b(log|audit trail|backup)\b'
        ]
        
        # Extract using patterns
        for pattern in entity_patterns:
            matches = re.findall(pattern, content_lower)
            for match in set(matches):
                clean_match = match.strip().title()
                if clean_match not in result.external_entities:
                    result.external_entities.append(clean_match)
        
        for pattern in process_patterns:
            matches = re.findall(pattern, content_lower)
            for match in set(matches):
                clean_match = match.strip().title()
                if clean_match not in result.processes:
                    result.processes.append(clean_match)
        
        for pattern in asset_patterns:
            matches = re.findall(pattern, content_lower)
            for match in set(matches):
                clean_match = match.strip().title()
                if clean_match not in result.assets:
                    result.assets.append(clean_match)
        
        # Add default components if none found
        if not result.external_entities:
            result.external_entities = ["User", "Administrator"]
        if not result.processes:
            result.processes = ["Web Server", "Application Server"]
        if not result.assets:
            result.assets = ["Database", "File Storage"]
        
        # Add basic trust boundaries
        result.trust_boundaries = ["External to Internal", "DMZ to Application", "Application to Data"]
        
        # Add basic data flows
        if len(result.external_entities) > 0 and len(result.processes) > 0:
            flow = SimpleDataFlow(
                source=result.external_entities[0],
                destination=result.processes[0],
                data_description="User requests and authentication data",
                data_classification="Internal",
                protocol="HTTPS",
                authentication_mechanism="Session Token"
            )
            result.data_flows.append(flow)
        
        if len(result.processes) > 0 and len(result.assets) > 0:
            flow = SimpleDataFlow(
                source=result.processes[0],
                destination=result.assets[0],
                data_description="Application data and user information",
                data_classification="Confidential",
                protocol="JDBC",
                authentication_mechanism="Database Credentials"
            )
            result.data_flows.append(flow)
        
        result.assumptions = ["Rule-based extraction used", "Limited detail available"]
        result.confidence_notes = ["Manual review recommended"]
        
        return result
    
    def _detect_industry(self, content: str) -> str:
        """Detect industry context from content."""
        content_lower = content.lower()
        
        if any(keyword in content_lower for keyword in ['payment', 'financial', 'bank', 'transaction', 'pci']):
            return "Financial"
        elif any(keyword in content_lower for keyword in ['patient', 'medical', 'health', 'hipaa', 'phi']):
            return "Healthcare"
        elif any(keyword in content_lower for keyword in ['ecommerce', 'retail', 'product', 'order', 'cart']):
            return "E-commerce"
        elif any(keyword in content_lower for keyword in ['saas', 'software', 'tenant', 'subscription']):
            return "SaaS"
        else:
            return "General"

# =====================================================================
# IMPROVED MERMAID DIAGRAM GENERATOR
# =====================================================================

def generate_threat_modeling_mermaid(dfd_data: Dict[str, Any]) -> str:
    """Generate a threat modeling focused Mermaid diagram."""
    logger.info("🎨 Generating threat modeling Mermaid diagram...")
    
    if not dfd_data or 'dfd' not in dfd_data:
        logger.warning("No DFD data available for Mermaid generation")
        return ""
    
    dfd = dfd_data['dfd']
    lines = ['graph TB']
    
    def safe_id(text: str) -> str:
        """Create safe Mermaid node ID."""
        if not text:
            return 'unknown'
        # Replace non-alphanumeric with underscore, limit length
        safe = re.sub(r'[^a-zA-Z0-9]', '_', str(text))
        safe = re.sub(r'_+', '_', safe).strip('_')
        if safe and safe[0].isdigit():
            safe = 'node_' + safe
        return safe[:20] or 'unknown'
    
    def get_trust_zone(name: str, comp_type: str) -> str:
        """Determine trust zone for component."""
        name_lower = name.lower()
        
        if comp_type == 'external' or any(keyword in name_lower for keyword in 
            ['user', 'client', 'external', 'internet', 'public']):
            return 'external'
        elif any(keyword in name_lower for keyword in 
            ['gateway', 'proxy', 'load balancer', 'firewall', 'waf']):
            return 'dmz'
        elif any(keyword in name_lower for keyword in 
            ['database', 'db', 'storage', 'cache', 'repository']):
            return 'data'
        else:
            return 'application'
    
    # Group components by trust zones
    zones = {'external': [], 'dmz': [], 'application': [], 'data': []}
    all_components = {}
    
    # Process components
    for entity in dfd.get('external_entities', []):
        zone = get_trust_zone(entity, 'external')
        comp_id = safe_id(entity)
        zones[zone].append({'id': comp_id, 'name': entity, 'type': 'entity'})
        all_components[entity] = comp_id
    
    for process in dfd.get('processes', []):
        zone = get_trust_zone(process, 'process')
        comp_id = safe_id(process)
        zones[zone].append({'id': comp_id, 'name': process, 'type': 'process'})
        all_components[process] = comp_id
    
    for asset in dfd.get('assets', []):
        zone = get_trust_zone(asset, 'asset')
        comp_id = safe_id(asset)
        zones[zone].append({'id': comp_id, 'name': asset, 'type': 'asset'})
        all_components[asset] = comp_id
    
    # Generate trust zone subgraphs
    zone_titles = {
        'external': '🌐 External Zone (Untrusted)',
        'dmz': '🛡️ DMZ Zone (Semi-Trusted)', 
        'application': '🏢 Application Zone (Trusted)',
        'data': '💾 Data Zone (Critical Assets)'
    }
    
    for zone, components in zones.items():
        if components:
            lines.append(f'    subgraph {zone}["{zone_titles[zone]}"]')
            for comp in components:
                icon = '👤' if comp['type'] == 'entity' else '💾' if comp['type'] == 'asset' else '⚙️'
                lines.append(f'        {comp["id"]}["{icon} {comp["name"]}"]')
            lines.append('    end')
            lines.append('')
    
    # Add data flows
    if dfd.get('data_flows'):
        lines.append('    %% Data Flows with Security Context')
        for flow in dfd['data_flows']:
            source_id = all_components.get(flow.get('source'))
            dest_id = all_components.get(flow.get('destination'))
            
            if not source_id or not dest_id:
                continue
            
            # Determine arrow style based on risk
            data_class = flow.get('data_classification', 'Internal')
            if data_class in ['PII', 'PHI', 'PCI']:
                arrow = '==>'  # High risk
            elif data_class == 'Confidential':
                arrow = '-->'  # Medium risk
            else:
                arrow = '-.->'  # Low risk
            
            # Create label
            protocol = flow.get('protocol', 'Unknown')
            auth = flow.get('authentication_mechanism', 'None')
            encrypted = '🔒' if flow.get('encryption_in_transit') else '🔓'
            
            label = f"{protocol}|{data_class}|{auth[:10]}|{encrypted}"
            lines.append(f'    {source_id} {arrow}|"{label}"| {dest_id}')
    
    # Add styling
    lines.extend([
        '',
        '    %% Trust Zone Styling',
        '    classDef external fill:#ff4757,stroke:#ff3742,stroke-width:3px,color:#fff',
        '    classDef dmz fill:#ffa502,stroke:#ff8c00,stroke-width:2px,color:#000',
        '    classDef application fill:#3742fa,stroke:#2f40fa,stroke-width:2px,color:#fff',
        '    classDef data fill:#2ed573,stroke:#20bf6b,stroke-width:2px,color:#000',
        ''
    ])
    
    # Apply zone classes
    for zone, components in zones.items():
        if components:
            comp_ids = [comp['id'] for comp in components]
            lines.append(f'    class {",".join(comp_ids)} {zone}')
    
    # Add legend
    lines.extend([
        '',
        '    %% THREAT MODELING LEGEND:',
        '    %% 🌐 Red External: Untrusted attack surface',
        '    %% 🛡️ Orange DMZ: Semi-trusted exposed services', 
        '    %% 🏢 Blue Application: Trusted business logic',
        '    %% 💾 Green Data: Critical assets needing protection',
        '    %% === High-risk data flows (PII/PHI/PCI)',
        '    %% --- Medium-risk flows (Confidential)',
        '    %% -.- Low-risk flows (Internal/Public)'
    ])
    
    result = '\n'.join(lines)
    logger.info(f"✅ Generated Mermaid diagram with {len(lines)} lines")
    return result

# =====================================================================
# MAIN EXTRACTOR CLASS
# =====================================================================

class ImprovedDFDExtractor:
    """Improved DFD extractor with better error handling and validation."""
    
    def __init__(self):
        self.llm_client = SimpleLLMClient(config['llm_provider'], config['llm_model'])
    
    def extract_dfd_from_documents(self, documents: List[str], document_info: List[str], step: int = 2) -> Dict[str, Any]:
        """Extract DFD from documents with improved processing."""
        
        logger.info("🚀 Starting improved DFD extraction")
        write_progress(step, 5, 100, "Starting extraction", "Analyzing documents")
        
        # Combine and analyze documents
        combined_content = "\n\n--- DOCUMENT SEPARATOR ---\n\n".join(documents)
        total_length = len(combined_content)
        
        logger.info(f"📝 Processing {total_length} characters from {len(documents)} documents")
        logger.info(f"📄 Documents: {document_info}")
        
        if check_kill_signal(step):
            return self._create_error_result("Extraction cancelled by user")
        
        # Document analysis
        write_progress(step, 15, 100, "Analyzing content", "Detecting structure and context")
        doc_analysis = self._analyze_document_content(combined_content)
        logger.info(f"📊 Analysis: {doc_analysis['document_type']} | {doc_analysis['industry_context']}")
        
        if check_kill_signal(step):
            return self._create_error_result("Extraction cancelled by user")
        
        # Extract components
        write_progress(step, 30, 100, "Extracting components", "Running AI analysis")
        extraction_result = self.llm_client.extract_dfd_components(combined_content, doc_analysis)
        
        if not extraction_result:
            logger.error("Component extraction failed")
            write_progress(step, 100, 100, "Failed", "Component extraction failed")
            return self._create_error_result("Component extraction failed")
        
        if check_kill_signal(step):
            return self._create_error_result("Extraction cancelled by user")
        
        # Validate and improve extraction
        write_progress(step, 60, 100, "Validating results", "Checking consistency")
        validation_results = self._validate_extraction(extraction_result)
        
        # Convert to dictionary format
        dfd_dict = extraction_result.to_dict()
        
        # Generate Mermaid diagram
        mermaid_diagram = ""
        if config['enable_mermaid']:
            write_progress(step, 80, 100, "Generating diagram", "Creating visualization")
            try:
                mermaid_diagram = generate_threat_modeling_mermaid({"dfd": dfd_dict})
                logger.info("🎨 Mermaid diagram generated successfully")
            except Exception as e:
                logger.warning(f"Mermaid generation failed: {e}")
        
        # Store global reference for frontend
        import __main__
        __main__.currentDfdData = {"dfd": dfd_dict}
        
        # Create final output
        write_progress(step, 95, 100, "Finalizing", "Creating comprehensive output")
        final_result = {
            "dfd": dfd_dict,
            "mermaid": mermaid_diagram,
            "metadata": {
                "timestamp": datetime.now().isoformat(),
                "extraction_version": "4.0_improved",
                "source_documents": document_info,
                "document_analysis": doc_analysis,
                "validation_results": validation_results,
                "llm_provider": config['llm_provider'],
                "llm_model": config['llm_model'],
                "total_content_length": total_length,
                "extraction_stats": {
                    "external_entities": len(dfd_dict.get('external_entities', [])),
                    "processes": len(dfd_dict.get('processes', [])),
                    "assets": len(dfd_dict.get('assets', [])),
                    "data_flows": len(dfd_dict.get('data_flows', [])),
                    "trust_boundaries": len(dfd_dict.get('trust_boundaries', []))
                },
                "quality_indicators": {
                    "has_trust_boundaries": len(dfd_dict.get('trust_boundaries', [])) > 0,
                    "has_data_classification": any(
                        flow.get('data_classification') not in ['Unknown', 'Internal'] 
                        for flow in dfd_dict.get('data_flows', [])
                    ),
                    "has_authentication": any(
                        flow.get('authentication_mechanism') not in ['Unknown', 'None']
                        for flow in dfd_dict.get('data_flows', [])
                    )
                }
            }
        }
        
        write_progress(step, 100, 100, "Complete", f"Extracted {self._count_components(dfd_dict)} components")
        logger.info("✅ DFD extraction completed successfully")
        
        return final_result
    
    def _analyze_document_content(self, content: str) -> Dict[str, Any]:
        """Analyze document content to guide extraction."""
        content_lower = content.lower()
        
        # Detect document type
        doc_type = "technical_requirements"
        if "architecture" in content_lower:
            doc_type = "architecture_document"
        elif "api" in content_lower and ("endpoint" in content_lower or "rest" in content_lower):
            doc_type = "api_documentation"
        elif "security" in content_lower and "threat" in content_lower:
            doc_type = "security_document"
        
        # Detect industry
        industry = "General"
        industry_keywords = {
            "Financial": ["payment", "transaction", "banking", "fintech", "pci", "fraud"],
            "Healthcare": ["patient", "medical", "hipaa", "phi", "healthcare", "clinical"],
            "E-commerce": ["cart", "checkout", "order", "product", "inventory", "customer"],
            "SaaS": ["tenant", "subscription", "api", "saas", "software"]
        }
        
        for industry_name, keywords in industry_keywords.items():
            if sum(1 for keyword in keywords if keyword in content_lower) >= 2:
                industry = industry_name
                break
        
        # Calculate complexity
        complexity = min(len(content) / 10000, 1.0)
        technical_terms = len(re.findall(r'\b(?:system|service|database|api|server|user|data|process)\b', content_lower))
        complexity += min(technical_terms / 50, 1.0)
        complexity = min(complexity / 2, 1.0)
        
        return {
            "document_type": doc_type,
            "industry_context": industry,
            "complexity_score": complexity,
            "content_length": len(content),
            "technical_term_count": technical_terms
        }
    
    def _validate_extraction(self, extraction: SimpleDFDComponents) -> Dict[str, Any]:
        """Validate extraction results."""
        validation = {
            "is_valid": True,
            "warnings": [],
            "errors": [],
            "completeness_score": 0.0
        }
        
        # Check minimum components
        total_components = (len(extraction.external_entities) + 
                          len(extraction.processes) + 
                          len(extraction.assets))
        
        if total_components < 3:
            validation["errors"].append(f"Insufficient components: {total_components} < 3")
            validation["is_valid"] = False
        
        # Check data flow consistency
        all_components = set(extraction.external_entities + extraction.processes + extraction.assets)
        
        for i, flow in enumerate(extraction.data_flows):
            if hasattr(flow, 'source'):
                source = flow.source
                dest = flow.destination
            else:
                source = flow.get('source', '')
                dest = flow.get('destination', '')
            
            if source not in all_components:
                validation["errors"].append(f"Data flow {i+1}: source '{source}' not in components")
            if dest not in all_components:
                validation["errors"].append(f"Data flow {i+1}: destination '{dest}' not in components")
        
        # Calculate completeness
        factors = {
            "has_external_entities": len(extraction.external_entities) > 0,
            "has_processes": len(extraction.processes) > 0,
            "has_assets": len(extraction.assets) > 0,
            "has_data_flows": len(extraction.data_flows) > 0,
            "has_trust_boundaries": len(extraction.trust_boundaries) > 0
        }
        
        validation["completeness_score"] = sum(factors.values()) / len(factors)
        
        logger.info(f"🔍 Validation: {len(validation['errors'])} errors, {len(validation['warnings'])} warnings")
        
        return validation
    
    def _count_components(self, dfd_dict: Dict) -> int:
        """Count total components."""
        return (len(dfd_dict.get('external_entities', [])) +
                len(dfd_dict.get('processes', [])) + 
                len(dfd_dict.get('assets', [])))
    
    def _create_error_result(self, error_message: str) -> Dict[str, Any]:
        """Create error result structure."""
        return {
            "dfd": {
                "project_name": "Error",
                "project_version": "1.0",
                "industry_context": "unknown",
                "external_entities": [],
                "processes": [],
                "assets": [],
                "trust_boundaries": [],
                "data_flows": [],
                "error": error_message
            },
            "mermaid": "",
            "metadata": {
                "timestamp": datetime.now().isoformat(),
                "error": error_message,
                "extraction_version": "4.0_improved",
                "status": "failed"
            }
        }

# =====================================================================
# MAIN EXECUTION
# =====================================================================

def main():
    """Main execution function with comprehensive improvements."""
    logger.info("\n🛡️ === Improved DFD Extraction Engine ===")
    write_progress(2, 0, 100, "Initializing", "Starting improved DFD extraction")

    try:
        # Load and validate documents
        documents, document_info = load_documents_from_step1(config['output_dir'], step=2)
        
        if not documents:
            raise ValueError("No valid documents loaded")
        
        if check_kill_signal(2):
            logger.info("🛑 Extraction cancelled by user")
            return 1
        
        # Initialize extractor and run extraction
        extractor = ImprovedDFDExtractor()
        result = extractor.extract_dfd_from_documents(documents, document_info, step=2)
        
        if check_kill_signal(2):
            logger.info("🛑 Extraction cancelled by user")
            return 1
        
        # Validate output structure
        if not result or 'dfd' not in result:
            raise ValueError("Invalid extraction result structure")
        
        # Save results
        with open(config['dfd_output_path'], 'w', encoding='utf-8') as f:
            json.dump(result, f, indent=2, ensure_ascii=False)
        
        logger.info(f"💾 Results saved to: {config['dfd_output_path']}")
        
        # Print summary
        print_extraction_summary(result)
        
        # Clean up progress file
        try:
            progress_file = os.path.join(config['output_dir'], 'step_2_progress.json')
            if os.path.exists(progress_file):
                os.remove(progress_file)
        except:
            pass
        
        return 0

    except Exception as e:
        logger.error(f"❌ DFD extraction failed: {e}")
        write_progress(2, 100, 100, "Failed", str(e))
        return 1

def print_extraction_summary(result: Dict[str, Any]):
    """Print comprehensive extraction summary."""
    dfd = result.get('dfd', {})
    metadata = result.get('metadata', {})
    
    print(f"\n{'='*80}")
    print("🛡️  IMPROVED DFD EXTRACTION SUMMARY")
    print(f"{'='*80}")
    
    # Component stats
    stats = metadata.get('extraction_stats', {})
    print(f"📊 Components Extracted:")
    print(f"   • External Entities: {stats.get('external_entities', 0)}")
    print(f"   • Processes: {stats.get('processes', 0)}")
    print(f"   • Assets: {stats.get('assets', 0)}")
    print(f"   • Data Flows: {stats.get('data_flows', 0)}")
    print(f"   • Trust Boundaries: {stats.get('trust_boundaries', 0)}")
    
    # Quality indicators
    quality = metadata.get('quality_indicators', {})
    print(f"\n🔍 Quality Indicators:")
    print(f"   • Trust Boundaries: {'✅' if quality.get('has_trust_boundaries') else '❌'}")
    print(f"   • Data Classification: {'✅' if quality.get('has_data_classification') else '❌'}")
    print(f"   • Authentication Details: {'✅' if quality.get('has_authentication') else '❌'}")
    
    # Document analysis
    doc_analysis = metadata.get('document_analysis', {})
    print(f"\n📄 Document Analysis:")
    print(f"   • Type: {doc_analysis.get('document_type', 'Unknown')}")
    print(f"   • Industry: {doc_analysis.get('industry_context', 'Unknown')}")
    print(f"   • Complexity: {doc_analysis.get('complexity_score', 0)*100:.1f}%")
    print(f"   • Content Length: {doc_analysis.get('content_length', 0):,} chars")
    
    # Validation results
    validation = metadata.get('validation_results', {})
    print(f"\n✅ Validation:")
    print(f"   • Valid: {'✅' if validation.get('is_valid', False) else '❌'}")
    print(f"   • Completeness: {validation.get('completeness_score', 0)*100:.1f}%")
    
    if validation.get('errors'):
        print(f"   • Errors: {len(validation['errors'])}")
    if validation.get('warnings'):
        print(f"   • Warnings: {len(validation['warnings'])}")
    
    # Mermaid diagram
    mermaid_available = bool(result.get('mermaid'))
    print(f"\n🎨 Visualization:")
    print(f"   • Mermaid Diagram: {'✅ Generated' if mermaid_available else '❌ Failed'}")
    
    # Output info
    print(f"\n📁 Output:")
    print(f"   • File: {config['dfd_output_path']}")
    print(f"   • Version: {metadata.get('extraction_version', 'Unknown')}")
    
    print(f"{'='*80}")
    
    # Show validation issues if any
    if validation.get('errors') or validation.get('warnings'):
        print("⚠️  VALIDATION ISSUES:")
        for error in validation.get('errors', []):
            print(f"   ❌ {error}")
        for warning in validation.get('warnings', []):
            print(f"   ⚠️  {warning}")
        print(f"{'='*80}")

if __name__ == "__main__":
    sys.exit(main())